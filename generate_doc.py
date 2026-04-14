from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()

    # Title
    title = doc.add_heading('Dokumentasi Alur Autentikasi Nyutji Mobile', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph('Dokumentasi ini menjelaskan alur kerja, desain, dan logika di balik sistem Login dan Registrasi pada aplikasi Nyutji Mobile.')

    # Section 1
    doc.add_heading('1. Ikhtisar Sistem', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('Sistem autentikasi dirancang dengan pendekatan ').font.bold = False
    p1.add_run('Retro-Premium').font.bold = True
    p1.add_run(' menggunakan Flutter. Fitur utama meliputi:').font.bold = False
    
    doc.add_paragraph('Global Localization: Dukungan penuh Bahasa Indonesia (ID) dan Inggris (EN).', style='List Bullet')
    doc.add_paragraph('Multi-Role Registration: Pendaftaran khusus untuk Pelanggan, Kurir, dan Mitra.', style='List Bullet')
    doc.add_paragraph('State Management: Menggunakan AuthProvider untuk sinkronisasi status login dan bahasa.', style='List Bullet')

    # Section 2
    doc.add_heading('2. Alur Login (LoginScreen)', level=1)
    doc.add_paragraph('Layar pertama yang dilihat pengguna saat membuka aplikasi.')
    doc.add_paragraph('Language Toggle: Tombol di pojok kanan atas untuk mengganti bahasa secara global.', style='List Bullet')
    doc.add_paragraph('Dua Tahap Verifikasi: Nomor HP (awal) dan Password (setelah verifikasi HP).', style='List Bullet')
    doc.add_paragraph('Navigasi: Tombol Daftar Sekarang mengarah ke RegisterScreen.', style='List Bullet')

    # Section 3
    doc.add_heading('3. Alur Registrasi Utama (RegisterScreen)', level=1)
    doc.add_paragraph('Role Cards: Tiga tombol bulat besar dengan ikon (Pelanggan, Kurir, Mitra).')
    doc.add_paragraph('Logika Navigasi: Auto-Swipe menuju layar spesifik untuk Kurir & Mitra.', style='List Bullet')
    doc.add_paragraph('Pelanggan: Form pendaftaran muncul di layar yang sama.', style='List Bullet')

    # Section 4
    doc.add_heading('4. Registrasi Kurir (RegisterKurirScreen)', level=1)
    doc.add_paragraph('Tema: Retro Orange.')
    doc.add_paragraph('Data: Nama, Email, HP, Password, Lokasi Penugasan (Map), Referensi Mitra.', style='List Bullet')

    # Section 5
    doc.add_heading('5. Registrasi Mitra (RegisterMitraScreen)', level=1)
    doc.add_paragraph('Tema: Retro Green.')
    doc.add_paragraph('Segment: Pribadi atau Badan Usaha (dengan ikon & deskripsi).', style='List Bullet')
    doc.add_paragraph('Kategori: Kecil, Sedang, Besar (dengan kapasitas harian).', style='List Bullet')
    doc.add_paragraph('Data: Nama Pemilik, Email, HP Bisnis, Password, Lokasi Wilayah.', style='List Bullet')

    # Section 6
    doc.add_heading('6. Teknis Implementasi', level=1)
    doc.add_paragraph('Routing: RetroRoute (Custom PageRouteBuilder) untuk transisi slide.', style='List Bullet')
    doc.add_paragraph('Color Palette: Pelanggan (#286B6A), Kurir (#D35400), Mitra (#27AE60).', style='List Bullet')

    # Save
    doc.save('Nyutji_Auth_Flow_Documentation.docx')
    print("Document created successfully!")

if __name__ == "__main__":
    create_doc()
